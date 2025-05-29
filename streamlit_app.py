import streamlit as st
import os
import zipfile
import shutil
tempfile
import urllib.parse
import hashlib
import requests
from io import BytesIO
from docx import Document

# ─── Streamlit Page Config (MUST be first) ───
st.set_page_config(page_title="SCORM Review Links")
st.title("SCORM Review Link Generator")

# ─── Backblaze B2 Credentials ───
B2_KEY_ID   = "0057d19983190740000000003"
B2_APP_KEY  = "K005tH7/zOTPfBmFlyDMy2cYamvw5y8"
BUCKET_NAME = "filesfornecym"
BASE_URL    = f"https://f005.backblazeb2.com/file/{BUCKET_NAME}"
API_VER     = "v2"

# ─── Helper: Authorize & Cache for 1h ───
@st.cache_data(ttl=3600)
def authorize_b2():
    resp = requests.get(
        f"https://api.backblazeb2.com/b2api/{API_VER}/b2_authorize_account",
        auth=(B2_KEY_ID, B2_APP_KEY)
    )
    resp.raise_for_status()
    return resp.json()

try:
    auth = authorize_b2()
    API_URL    = auth["apiUrl"]
    AUTH_TOKEN = auth["authorizationToken"]
    ACCOUNT_ID = auth["accountId"]
except Exception as e:
    st.error(f"❌ Authorization failed: {e}")
    st.stop()

# ─── Helper: Get Bucket ID ───
@st.cache_data(ttl=3600)
def get_bucket_id():
    payload = {"accountId": ACCOUNT_ID}
    resp = requests.post(
        f"{API_URL}/b2api/{API_VER}/b2_list_buckets",
        headers={"Authorization": AUTH_TOKEN},
        json=payload
    )
    resp.raise_for_status()
    for b in resp.json()["buckets"]:
        if b["bucketName"] == BUCKET_NAME:
            return b["bucketId"]
    raise RuntimeError(f"Bucket {BUCKET_NAME!r} not found")

try:
    BUCKET_ID = get_bucket_id()
    st.success("✅ Connected to Backblaze B2.")
except Exception as e:
    st.error(f"❌ Failed to find bucket: {e}")
    st.stop()

# ─── B2 Operations ───
def list_files(prefix=""):
    payload = {"bucketId": BUCKET_ID, "prefix": prefix, "maxFileCount": 10000}
    resp = requests.post(
        f"{API_URL}/b2api/{API_VER}/b2_list_file_names",
        headers={"Authorization": AUTH_TOKEN},
        json=payload
    )
    resp.raise_for_status()
    return resp.json()["files"]  # each entry has fileName, fileId

def delete_file_version(file_name, file_id):
    payload = {"fileName": file_name, "fileId": file_id}
    resp = requests.post(
        f"{API_URL}/b2api/{API_VER}/b2_delete_file_version",
        headers={"Authorization": AUTH_TOKEN},
        json=payload
    )
    resp.raise_for_status()

def get_upload_url():
    payload = {"bucketId": BUCKET_ID}
    resp = requests.post(
        f"{API_URL}/b2api/{API_VER}/b2_get_upload_url",
        headers={"Authorization": AUTH_TOKEN},
        json=payload
    )
    resp.raise_for_status()
    j = resp.json()
    return j["uploadUrl"], j["authorizationToken"]

def upload_file(local_path, b2_path):
    upload_url, upload_token = get_upload_url()
    with open(local_path, "rb") as f:
        data = f.read()
    sha1 = hashlib.sha1(data).hexdigest()
    headers = {
        "Authorization": upload_token,
        "X-Bz-File-Name": urllib.parse.quote(b2_path, safe=""),
        "Content-Type": "b2/x-auto",
        "X-Bz-Content-Sha1": sha1
    }
    resp = requests.post(upload_url, headers=headers, data=data)
    resp.raise_for_status()

# ─── Build Word Document of Links ───
def build_links_doc(langs, files):
    doc = Document()
    doc.add_heading("SCORM Review Links", 0)

    for lang in sorted(langs):
        doc.add_heading(lang, level=1)
        packages = {
            fn.split("/")[1]
            for fn in (f["fileName"] for f in files)
            if fn.startswith(f"{lang}/") and fn.count("/") >= 2
        }
        if not packages:
            doc.add_paragraph("*(no packages found)*")
            continue

        for pkg in sorted(packages):
            url = f"{BASE_URL}/{urllib.parse.quote(lang)}/{urllib.parse.quote(pkg)}/story.html"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{pkg}: ")
            run = p.add_run(url)
            run.font.color.theme_color = 10
            run.font.underline = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ─── Streamlit UI ───
mode = st.sidebar.radio("Mode", ["View Links", "Upload New Language"])

if mode == "View Links":
    files = list_files("")
    langs = {f["fileName"].split("/",1)[0] for f in files if "/" in f["fileName"]}
    selected = st.selectbox("Select Language:", sorted(langs))

    if selected:
        pkgs = {
            fn.split("/")[1]
            for fn in (f["fileName"] for f in files)
            if fn.startswith(f"{selected}/") and fn.count("/") >= 2
        }
        if pkgs:
            st.subheader(f"SCORM Packages in {selected}:")
            for pkg in sorted(pkgs):
                safe = pkg.replace("[","\\[").replace("]","\\]")
                link = f"{BASE_URL}/{urllib.parse.quote(selected)}/{urllib.parse.quote(pkg)}/story.html"
                st.markdown(f"📄 [{safe}]({link})")
        else:
            st.info("No SCORM packages found.")

        # Export to Word
        if st.button("📄 Export all links to Word"):
            docx_buf = build_links_doc(langs, files)
            st.download_button(
                label="⬇️ Download SCORM Links .docx",
                data=docx_buf,
                file_name="scorm_links.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

elif mode == "Upload New Language":
    existing = {f["fileName"].split("/",1)[0] for f in list_files("") if "/" in f["fileName"]}
    lang = st.text_input("Language folder name (e.g. German):")

    if lang:
        if lang in existing:
            if st.checkbox(f"⚠️ Delete existing '{lang}/' first"):
                if st.button("Delete & Start Upload"):
                    with st.spinner("Deleting…"):
                        for f in list_files(f"{lang}/"):
                            delete_file_version(f["fileName"], f["fileId"])
                    st.success("✅ Deleted old files.")
        else:
            st.info(f"'{lang}' is new; ready to upload.")

        uploads = st.file_uploader("SCORM ZIPs", type="zip", accept_multiple_files=True)
        if uploads and st.button("Start Upload"):
            with st.spinner("Uploading…"):
                for z in uploads:
                    base = os.path.splitext(z.name)[0]
                    tmp = tempfile.mkdtemp()
                    path = os.path.join(tmp, z.name)
                    with open(path,"wb") as f: f.write(z.read())
                    ext = os.path.join(tmp, base)
                    os.makedirs(ext, exist_ok=True)
                    with zipfile.ZipFile(path) as zp:
                        zp.extractall(ext)
                    for root,_,fs in os.walk(ext):
                        for fn in fs:
                            lp = os.path.join(root, fn)
                            rel = os.path.relpath(lp, ext).replace("\\","/")
                            b2p = f"{lang}/{base}/{rel}"
                            upload_file(lp, b2p)
                    shutil.rmtree(tmp)
                st.success("✅ Upload complete.")

st.caption("Developed for instant SCORM review link generation and management.")
